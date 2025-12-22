"""DOCX Renderer for editable human-readable output

Renders canonical JSON to editable Microsoft Word format.
Follows LLM_TECHNICAL_SPEC.md Section 7: Human-Readable Output

CRITICAL: This is a RENDERER ONLY - generates output from canonical JSON.
NO data extraction logic allowed here.
"""

from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models.canonical_schema import MedicalDocument
from ..utils.logger import get_logger

logger = get_logger(__name__)


class DOCXRenderer:
    """Render canonical JSON to editable DOCX document"""

    def __init__(self):
        logger.info("DOCX renderer initialized")

    def render(self, document: MedicalDocument, output_path: str) -> str:
        """Render MedicalDocument to DOCX file

        Args:
            document: Validated MedicalDocument instance
            output_path: Path to save DOCX file

        Returns:
            Path to saved DOCX file
        """
        logger.info("Rendering document to DOCX", visits=len(document.visits))

        try:
            # Create document
            doc = Document()

            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            # Add header
            self._add_header(doc, document)

            # Add disclaimer
            self._add_disclaimer(doc)

            # Add patient demographics
            self._add_patient_demographics(doc, document)

            # Add each visit
            for i, visit in enumerate(document.visits):
                if i > 0:
                    doc.add_page_break()
                self._add_visit(doc, visit, i + 1)

            # Add data quality section
            self._add_data_quality(doc, document)

            # Save document
            doc.save(output_path)

            logger.info("DOCX rendering complete", output_path=output_path)
            return output_path

        except Exception as e:
            logger.error("DOCX rendering failed", error=str(e))
            raise RenderError(f"Failed to render DOCX: {e}")

    def _add_header(self, doc: Document, document: MedicalDocument):
        """Add document header"""
        title = doc.add_heading("MEDICAL RECORD SUMMARY", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'

        headers = ["Patient:", "DOB:", "Sex:", "Document Type:", "Generated:", "Source:"]
        values = [
            document.document_metadata.patient_name or "Unknown",
            str(document.document_metadata.dob) if document.document_metadata.dob else "Unknown",
            document.document_metadata.sex or "Unknown",
            document.document_metadata.document_type or "Mixed",
            document.document_metadata.processing_metadata.processed_at.strftime("%Y-%m-%d %H:%M:%S"),
            "OCR-processed medical record",
        ]

        for i, (header, value) in enumerate(zip(headers, values)):
            row = table.rows[i]
            row.cells[0].text = header
            row.cells[1].text = value

            # Bold headers
            run = row.cells[0].paragraphs[0].runs[0]
            run.bold = True

        doc.add_paragraph()

    def _add_disclaimer(self, doc: Document):
        """Add important disclaimer"""
        p = doc.add_paragraph()
        p.add_run("IMPORTANT: ").bold = True
        disclaimer = p.add_run(
            "This document was generated from OCR-processed scanned medical records. "
            "All information should be reviewed by qualified healthcare professionals. "
            "Do not use as the sole source of patient information for clinical decision-making. "
            "Always refer to original source documents for critical decisions."
        )
        # Red color for warning
        font = disclaimer.font
        font.color.rgb = RGBColor(231, 76, 60)

        doc.add_paragraph()

    def _add_patient_demographics(self, doc: Document, document: MedicalDocument):
        """Add patient demographics section"""
        doc.add_heading("PATIENT DEMOGRAPHICS", 1)

        demo_items = []
        if document.document_metadata.patient_name:
            demo_items.append(("Name:", document.document_metadata.patient_name))
        if document.document_metadata.patient_id:
            demo_items.append(("Patient ID:", str(document.document_metadata.patient_id)))
        if document.document_metadata.dob:
            demo_items.append(("Date of Birth:", str(document.document_metadata.dob)))
        if document.document_metadata.sex:
            demo_items.append(("Sex:", document.document_metadata.sex))

        if demo_items:
            for label, value in demo_items:
                p = doc.add_paragraph()
                p.add_run(f"{label} ").bold = True
                p.add_run(value)
        else:
            doc.add_paragraph("No patient demographics available")

        doc.add_paragraph()

    def _add_visit(self, doc: Document, visit: Dict[str, Any], visit_number: int):
        """Add a visit section"""
        # Visit header
        visit_title = f"VISIT {visit_number}: {visit.get('visit_date', 'Unknown Date')}"
        doc.add_heading(visit_title, 1)

        # Visit metadata
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light List Accent 1'

        metadata = [
            ("Visit ID:", visit.get("visit_id", "N/A")),
            ("Visit Date:", visit.get("visit_date", "N/A")),
            ("Encounter Type:", visit.get("encounter_type", "N/A")),
            ("Source Pages:", ", ".join(map(str, visit.get("raw_source_pages", [])))),
        ]

        for i, (label, value) in enumerate(metadata):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Clinical sections
        self._add_text_section(doc, "Reason for Visit", visit.get("reason_for_visit"))
        self._add_text_section(doc, "History of Present Illness", visit.get("history_of_present_illness"))
        self._add_medications_section(doc, visit.get("medications", []))
        self._add_vital_signs_section(doc, visit.get("vital_signs", {}))
        self._add_problem_list_section(doc, visit.get("problem_list", []))
        self._add_results_section(doc, visit.get("results", []))
        self._add_text_section(doc, "Assessment", visit.get("assessment"))
        self._add_plan_section(doc, visit.get("plan", []))

        # Review flags
        if visit.get("manual_review_required"):
            p = doc.add_paragraph()
            p.add_run("⚠ MANUAL REVIEW REQUIRED: ").bold = True
            warning_text = ", ".join(visit.get("review_reasons", []))
            run = p.add_run(warning_text)
            run.font.color.rgb = RGBColor(231, 76, 60)

    def _add_text_section(self, doc: Document, title: str, content: str):
        """Add a simple text section"""
        if not content:
            return

        doc.add_heading(title.upper(), 2)
        doc.add_paragraph(content)

    def _add_medications_section(self, doc: Document, medications: List[Dict[str, Any]]):
        """Add medications section with table"""
        if not medications:
            return

        doc.add_heading("MEDICATIONS", 2)

        # Create table
        table = doc.add_table(rows=len(medications) + 1, cols=5)
        table.style = 'Light Grid Accent 1'

        # Header row
        headers = ["Medication", "Dose", "Frequency", "Route", "Page"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, med in enumerate(medications, start=1):
            row = table.rows[i]
            row.cells[0].text = med.get("name", "")
            row.cells[1].text = med.get("dose", "") or "N/A"
            row.cells[2].text = med.get("frequency", "") or "N/A"
            row.cells[3].text = med.get("route", "") or "N/A"
            row.cells[4].text = str(med.get("source_page", "")) or "N/A"

        doc.add_paragraph()

    def _add_vital_signs_section(self, doc: Document, vital_signs: Dict[str, Any]):
        """Add vital signs section"""
        if not vital_signs or not any(v.get("value") for v in vital_signs.values() if isinstance(v, dict)):
            return

        doc.add_heading("VITAL SIGNS", 2)

        # Collect vital sign data
        vitals_data = []
        for vital_name, vital_data in vital_signs.items():
            if isinstance(vital_data, dict) and vital_data.get("value"):
                vitals_data.append((
                    vital_name.replace("_", " ").title(),
                    str(vital_data["value"]),
                    vital_data.get("unit", "")
                ))

        if vitals_data:
            table = doc.add_table(rows=len(vitals_data) + 1, cols=3)
            table.style = 'Light Grid Accent 1'

            # Header
            headers = ["Vital Sign", "Value", "Unit"]
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            # Data
            for i, (name, value, unit) in enumerate(vitals_data, start=1):
                row = table.rows[i]
                row.cells[0].text = name
                row.cells[1].text = value
                row.cells[2].text = unit

        doc.add_paragraph()

    def _add_problem_list_section(self, doc: Document, problems: List[Dict[str, Any]]):
        """Add problem list section"""
        if not problems:
            return

        doc.add_heading("PROBLEM LIST", 2)

        for problem in problems:
            problem_text = f"• {problem.get('problem', '')}"
            if problem.get("icd10_code"):
                problem_text += f" (ICD-10: {problem['icd10_code']})"
            if problem.get("status"):
                problem_text += f" - {problem['status']}"
            if problem.get("source_page"):
                problem_text += f" [Page {problem['source_page']}]"

            doc.add_paragraph(problem_text)

        doc.add_paragraph()

    def _add_results_section(self, doc: Document, results: List[Dict[str, Any]]):
        """Add lab results section with table"""
        if not results:
            return

        doc.add_heading("LAB RESULTS", 2)

        # Create table
        table = doc.add_table(rows=len(results) + 1, cols=6)
        table.style = 'Light Grid Accent 1'

        # Header row
        headers = ["Test", "Value", "Unit", "Reference Range", "Flag", "Page"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, result in enumerate(results, start=1):
            row = table.rows[i]
            row.cells[0].text = result.get("test_name", "")
            row.cells[1].text = str(result.get("value", "")) or "N/A"
            row.cells[2].text = result.get("unit", "") or "N/A"
            row.cells[3].text = result.get("reference_range", "") or "N/A"
            row.cells[4].text = result.get("abnormal_flag", "") or "normal"
            row.cells[5].text = str(result.get("source_page", "")) or "N/A"

        doc.add_paragraph()

    def _add_plan_section(self, doc: Document, plan: List[Dict[str, Any]]):
        """Add plan of care section"""
        if not plan:
            return

        doc.add_heading("PLAN OF CARE", 2)

        for item in plan:
            plan_text = f"• {item.get('action', '')}"
            if item.get("category"):
                plan_text += f" ({item['category']})"
            if item.get("source_page"):
                plan_text += f" [Page {item['source_page']}]"

            doc.add_paragraph(plan_text)

        doc.add_paragraph()

    def _add_data_quality(self, doc: Document, document: MedicalDocument):
        """Add data quality notes"""
        doc.add_page_break()
        doc.add_heading("DATA QUALITY NOTES", 1)

        # Quality metrics
        p = doc.add_paragraph()
        p.add_run("OCR Confidence (Average): ").bold = True
        p.add_run(f"{document.document_metadata.processing_metadata.ocr_confidence_avg * 100:.1f}%\n")

        p.add_run("Pages Processed: ").bold = True
        p.add_run(f"{document.document_metadata.processing_metadata.page_count}\n")

        p.add_run("Processing Duration: ").bold = True
        p.add_run(f"{document.document_metadata.processing_metadata.processing_duration_ms / 1000:.2f} seconds\n")

        doc.add_paragraph()

        # Warnings
        if document.data_quality.unclear_sections:
            doc.add_heading("⚠ Sections with Reduced Confidence", 2)
            for unclear in document.data_quality.unclear_sections:
                warning_text = f"• {unclear.get('section', 'Unknown')} (Page {unclear.get('page', 'N/A')}): {unclear.get('reason', '')}"
                p = doc.add_paragraph(warning_text)
                p.runs[0].font.color.rgb = RGBColor(231, 76, 60)


class RenderError(Exception):
    """Raised when rendering fails"""
    pass
